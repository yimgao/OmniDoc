"""
Format Converter Agent
Converts documentation between different formats (Markdown, HTML, PDF, DOCX)
"""
from typing import Optional, List
from datetime import datetime
from pathlib import Path
from src.agents.base_agent import BaseAgent
from src.utils.file_manager import FileManager
from src.context.context_manager import ContextManager
from src.context.shared_context import AgentType, DocumentStatus, AgentOutput
from src.rate_limit.queue_manager import RequestQueue


class FormatConverterAgent(BaseAgent):
    """
    Format Converter Agent
    
    Converts documentation between formats:
    - Markdown to HTML
    - Markdown to PDF (via HTML)
    - Markdown to DOCX
    - Preserves formatting and structure
    """
    
    def __init__(
        self,
        provider_name: Optional[str] = None,
        model_name: Optional[str] = None,
        rate_limiter: Optional[RequestQueue] = None,
        file_manager: Optional[FileManager] = None,
        api_key: Optional[str] = None,
        **provider_kwargs
    ):
        """Initialize Format Converter Agent"""
        # FormatConverter doesn't actually use LLM, but inherits for consistency
        super().__init__(
            provider_name=provider_name,
            model_name=model_name,
            rate_limiter=rate_limiter,
            api_key=api_key,
            **provider_kwargs
        )
        
        self.file_manager = file_manager or FileManager(base_dir="docs/formats")
        self.supported_formats = ["html", "pdf", "docx"]
    
    def generate(self, markdown_content: str) -> str:
        """
        Generate method required by BaseAgent interface
        
        Args:
            markdown_content: Markdown content (for consistency with other agents)
        
        Returns:
            HTML representation of the markdown (default conversion)
        """
        return self.markdown_to_html(markdown_content)
    
    def markdown_to_html(self, markdown_content: str) -> str:
        """
        Convert Markdown to HTML
        
        Args:
            markdown_content: Markdown content to convert
        
        Returns:
            HTML content
        """
        try:
            import markdown
            html_content = markdown.markdown(
                markdown_content,
                extensions=['extra', 'codehilite', 'tables']
            )
            # Wrap in proper HTML structure with enhanced styling
            full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Documentation</title>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Helvetica Neue', Arial, sans-serif;
            max-width: 900px;
            margin: 0 auto;
            padding: 40px 20px;
            line-height: 1.8;
            color: #333;
            background-color: #fff;
        }}
        h1 {{ font-size: 2.5em; margin: 1em 0 0.5em; color: #2c3e50; border-bottom: 3px solid #3498db; padding-bottom: 0.3em; }}
        h2 {{ font-size: 2em; margin: 1.5em 0 0.5em; color: #34495e; border-bottom: 2px solid #95a5a6; padding-bottom: 0.3em; }}
        h3 {{ font-size: 1.5em; margin: 1.2em 0 0.5em; color: #555; }}
        h4, h5, h6 {{ margin: 1em 0 0.5em; color: #666; }}
        p {{ margin: 1em 0; text-align: justify; }}
        code {{
            background-color: #f8f9fa;
            padding: 2px 6px;
            border-radius: 4px;
            font-family: 'Monaco', 'Courier New', monospace;
            font-size: 0.9em;
            color: #e83e8c;
        }}
        pre {{
            background-color: #f8f9fa;
            padding: 16px;
            border-radius: 6px;
            overflow-x: auto;
            border-left: 4px solid #3498db;
            margin: 1em 0;
        }}
        pre code {{ background-color: transparent; padding: 0; color: #333; }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 1.5em 0;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }}
        th, td {{
            border: 1px solid #dee2e6;
            padding: 12px;
            text-align: left;
        }}
        th {{
            background-color: #3498db;
            color: white;
            font-weight: 600;
        }}
        tr:nth-child(even) {{ background-color: #f8f9fa; }}
        ul, ol {{ margin: 1em 0; padding-left: 2em; }}
        li {{ margin: 0.5em 0; }}
        blockquote {{
            border-left: 4px solid #3498db;
            padding-left: 1em;
            margin: 1em 0;
            color: #666;
            font-style: italic;
        }}
        a {{ color: #3498db; text-decoration: none; }}
        a:hover {{ text-decoration: underline; }}
        @media print {{
            body {{ max-width: 100%; padding: 0; }}
            h1, h2, h3 {{ page-break-after: avoid; }}
        }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""
            return full_html
        except ImportError:
            # Fallback: basic conversion without markdown library
            print("⚠️  Warning: markdown library not installed, using basic HTML conversion")
            html_content = markdown_content.replace('\n', '<br>\n')
            return f"<html><body>{html_content}</body></html>"
    
    def html_to_pdf(self, html_content: str, output_path: Optional[str] = None) -> str:
        """
        Convert HTML to PDF with enhanced styling
        
        Args:
            html_content: HTML content to convert
            output_path: Optional output file path
        
        Returns:
            Path to generated PDF file
        """
        # Suppress stderr BEFORE importing weasyprint (it prints during import)
        import sys
        import io
        import warnings
        
        stderr_backup = sys.stderr
        stderr_capture = io.StringIO()
        sys.stderr = stderr_capture
        
        # Also suppress warnings
        warnings.filterwarnings("ignore")
        
        try:
            # Import weasyprint with stderr suppressed
            from weasyprint import HTML, CSS
                    
                    if not output_path:
                        output_path = "documentation.pdf"
                    
                    # Ensure .pdf extension
                    if not output_path.endswith('.pdf'):
                        output_path = str(Path(output_path).with_suffix('.pdf'))
                    
                    # Enhanced CSS for PDF
                    pdf_css = CSS(string='''
                        @page {
                            size: A4;
                            margin: 2cm;
                            @top-center {
                                content: "Documentation";
                            }
                            @bottom-center {
                                content: "Page " counter(page) " of " counter(pages);
                            }
                        }
                        body {
                            font-size: 11pt;
                            line-height: 1.6;
                        }
                        h1 {
                            page-break-after: avoid;
                            margin-top: 1.5em;
                        }
                        h2 {
                            page-break-after: avoid;
                            margin-top: 1.2em;
                        }
                        pre {
                            page-break-inside: avoid;
                        }
                        table {
                            page-break-inside: avoid;
                        }
                    ''')
            
            if not output_path:
                output_path = "documentation.pdf"
            
            # Ensure .pdf extension
            if not output_path.endswith('.pdf'):
                output_path = str(Path(output_path).with_suffix('.pdf'))
            
            # Enhanced CSS for PDF
            pdf_css = CSS(string='''
                @page {
                    size: A4;
                    margin: 2cm;
                    @top-center {
                        content: "Documentation";
                    }
                    @bottom-center {
                        content: "Page " counter(page) " of " counter(pages);
                    }
                }
                body {
                    font-size: 11pt;
                    line-height: 1.6;
                }
                h1 {
                    page-break-after: avoid;
                    margin-top: 1.5em;
                }
                h2 {
                    page-break-after: avoid;
                    margin-top: 1.2em;
                }
                pre {
                    page-break-inside: avoid;
                }
                table {
                    page-break-inside: avoid;
                }
            ''')
            
            html_obj = HTML(string=html_content)
            pdf_path = self.file_manager.base_dir / output_path
            html_obj.write_pdf(pdf_path, stylesheets=[pdf_css])
            
            # Restore stderr AFTER successful conversion
            sys.stderr = stderr_backup
            
            return str(pdf_path.absolute())
            
        except (OSError, ImportError, Exception) as e:
            # Restore stderr before checking error
            sys.stderr = stderr_backup
            # Clear the captured stderr (don't print the WeasyPrint warnings)
            stderr_capture.seek(0)
            stderr_capture.truncate(0)
            
            # If it's a library loading error (macOS), skip PDF conversion gracefully
            error_str = str(e).lower()
            if "libgobject" in error_str or "dlopen" in error_str or "cannot load library" in error_str:
                raise ImportError(
                    "PDF conversion unavailable: System libraries not available. "
                    "PDF conversion requires additional system libraries on macOS. "
                    "HTML and DOCX formats are still available."
                )
            raise
        except ImportError:
            try:
                # Fallback to pdfkit if available
                import pdfkit
                
                if not output_path:
                    output_path = "documentation.pdf"
                if not output_path.endswith('.pdf'):
                    output_path = str(Path(output_path).with_suffix('.pdf'))
                
                pdf_path = self.file_manager.base_dir / output_path
                options = {
                    'page-size': 'A4',
                    'margin-top': '0.75in',
                    'margin-right': '0.75in',
                    'margin-bottom': '0.75in',
                    'margin-left': '0.75in',
                    'encoding': "UTF-8",
                    'no-outline': None
                }
                pdfkit.from_string(html_content, str(pdf_path), options=options)
                
                return str(pdf_path.absolute())
            except ImportError:
                raise ImportError(
                    "PDF conversion requires 'weasyprint' or 'pdfkit'. "
                    "Install with: pip install weasyprint"
                )
    
    def markdown_to_docx(self, markdown_content: str, output_path: Optional[str] = None) -> str:
        """
        Convert Markdown to DOCX
        
        Args:
            markdown_content: Markdown content to convert
            output_path: Optional output file path
        
        Returns:
            Path to generated DOCX file
        """
        try:
            from docx import Document
            from docx.shared import Inches
            import re
            
            if not output_path:
                output_path = "documentation.docx"
            if not output_path.endswith('.docx'):
                output_path = str(Path(output_path).with_suffix('.docx'))
            
            doc = Document()
            
            # Parse markdown and convert to docx
            lines = markdown_content.split('\n')
            for line in lines:
                # Headings
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                # Lists
                elif line.strip().startswith('- ') or line.strip().startswith('* '):
                    doc.add_paragraph(line.strip()[2:], style='List Bullet')
                # Regular paragraphs
                elif line.strip():
                    doc.add_paragraph(line.strip())
                else:
                    doc.add_paragraph()
            
            docx_path = self.file_manager.base_dir / output_path
            doc.save(str(docx_path))
            
            return str(docx_path.absolute())
        except ImportError:
            raise ImportError(
                "DOCX conversion requires 'python-docx'. "
                "Install with: pip install python-docx"
            )
    
    def convert(
        self,
        markdown_content: str,
        output_format: str,
        output_filename: Optional[str] = None
    ) -> str:
        """
        Convert Markdown content to specified format
        
        Args:
            markdown_content: Markdown content to convert
            output_format: Target format ('html', 'pdf', 'docx')
            output_filename: Optional output filename
        
        Returns:
            Path to converted file
        """
        if output_format.lower() not in self.supported_formats:
            raise ValueError(
                f"Unsupported format: {output_format}. "
                f"Supported formats: {', '.join(self.supported_formats)}"
            )
        
        if output_format.lower() == 'html':
            html_content = self.markdown_to_html(markdown_content)
            if not output_filename:
                output_filename = "documentation.html"
            file_path = self.file_manager.write_file(output_filename, html_content)
            return file_path
        
        elif output_format.lower() == 'pdf':
            html_content = self.markdown_to_html(markdown_content)
            return self.html_to_pdf(html_content, output_filename)
        
        elif output_format.lower() == 'docx':
            return self.markdown_to_docx(markdown_content, output_filename)
        
        else:
            raise ValueError(f"Format conversion not implemented: {output_format}")
    
    def convert_all_documents(
        self,
        documents: dict,
        formats: List[str],
        project_id: Optional[str] = None,
        context_manager: Optional[ContextManager] = None
    ) -> dict:
        """
        Convert all documents to multiple formats
        
        Args:
            documents: Dict mapping document names to markdown content
            formats: List of target formats (e.g., ['html', 'pdf'])
            project_id: Optional project ID for context
            context_manager: Optional context manager
        
        Returns:
            Dict mapping document names to converted file paths
        """
        results = {}
        
        print(f"🔄 Converting {len(documents)} documents to formats: {', '.join(formats)}")
        
        for doc_name, markdown_content in documents.items():
            doc_results = {}
            
            for fmt in formats:
                try:
                    # Extract base name from doc_name (handle both string and Path-like)
                    if doc_name:
                        base_name = str(Path(doc_name).stem) if '.' in str(doc_name) else str(doc_name)
                    else:
                        base_name = "document"
                    output_filename = f"{base_name}.{fmt}"
                    
                    file_path = self.convert(
                        markdown_content=markdown_content,
                        output_format=fmt,
                        output_filename=output_filename
                    )
                    
                    doc_results[fmt] = file_path
                    print(f"  ✅ Converted {doc_name} to {fmt}")
                    
                except Exception as e:
                    print(f"  ❌ Error converting {doc_name} to {fmt}: {e}")
                    doc_results[fmt] = None
            
            results[doc_name] = doc_results
        
        # Save to context if available
        if project_id and context_manager:
            output = AgentOutput(
                agent_type=AgentType.FORMAT_CONVERTER,
                document_type="format_conversions",
                content=str(results),  # JSON-like string of results
                file_path="",  # Multiple files, no single path
                status=DocumentStatus.COMPLETE,
                generated_at=datetime.now()
            )
            context_manager.save_agent_output(project_id, output)
            print(f"✅ Format conversions saved to shared context (project: {project_id})")
        
        return results

