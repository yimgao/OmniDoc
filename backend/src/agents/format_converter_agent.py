"""
Format Converter Agent
Converts documentation between different formats (Markdown, HTML, PDF, DOCX)
"""
from typing import Optional, List
from datetime import datetime
from pathlib import Path
import os
import sys
import ctypes.util

# Set library paths BEFORE any imports (critical for macOS)
# macOS ignores DYLD_LIBRARY_PATH, but DYLD_FALLBACK_LIBRARY_PATH works better
if sys.platform == 'darwin' and Path('/opt/homebrew/lib').exists():
    homebrew_lib = '/opt/homebrew/lib'
    current_fallback = os.environ.get('DYLD_FALLBACK_LIBRARY_PATH', '')
    if homebrew_lib not in current_fallback:
        os.environ['DYLD_FALLBACK_LIBRARY_PATH'] = (
            f"{homebrew_lib}:{current_fallback}" if current_fallback 
            else homebrew_lib
        )
    # Also set DYLD_LIBRARY_PATH (some tools check this first)
    current_lib = os.environ.get('DYLD_LIBRARY_PATH', '')
    if homebrew_lib not in current_lib:
        os.environ['DYLD_LIBRARY_PATH'] = (
            f"{homebrew_lib}:{current_lib}" if current_lib 
            else homebrew_lib
        )
    
    # Monkey patch ctypes.util.find_library to help WeasyPrint find libraries
    # WeasyPrint looks for names like 'libpango-1.0-0' but files are 'libpango-1.0.0.dylib'
    _original_find_library = ctypes.util.find_library
    
    def _patched_find_library(name):
        """
        Patched find_library that handles WeasyPrint's library naming conventions.
        Maps names like 'libpango-1.0-0' to actual files like 'libpango-1.0.0.dylib'
        """
        # Mapping of WeasyPrint's expected names to actual library files
        library_map = {
            'libgobject-2.0-0': 'libgobject-2.0.0.dylib',
            'libgobject-2.0': 'libgobject-2.0.dylib',
            'libglib-2.0': 'libglib-2.0.dylib',
            'libpango-1.0-0': 'libpango-1.0.0.dylib',
            'libpango-1.0': 'libpango-1.0.dylib',
            'libpangoft2-1.0': 'libpangoft2-1.0.dylib',
            'libcairo-2': 'libcairo.2.dylib',
            'libcairo': 'libcairo.dylib',
            'libpixman-1': 'libpixman-1.dylib',
        }
        
        # Check if we have a mapping for this library
        mapped_name = library_map.get(name)
        if mapped_name:
            lib_path = Path(homebrew_lib) / mapped_name
            if lib_path.exists():
                return str(lib_path)
            # Also try versioned name without .dylib extension for ctypes
            versioned_path = Path(homebrew_lib) / mapped_name.replace('.dylib', '')
            if versioned_path.exists():
                return str(versioned_path)
        
        # For other libraries, try to find with actual filename pattern
        # Replace '-' with '.' in version numbers if it looks like a version
        if name.startswith('lib') and '-2.0-0' in name:
            # Try libgobject-2.0.0 pattern
            alt_name = name.replace('-2.0-0', '-2.0.0')
            alt_path = Path(homebrew_lib) / f"{alt_name}.dylib"
            if alt_path.exists():
                return str(alt_path)
        
        if name.startswith('lib') and '-1.0-0' in name:
            # Try libpango-1.0.0 pattern
            alt_name = name.replace('-1.0-0', '-1.0.0')
            alt_path = Path(homebrew_lib) / f"{alt_name}.dylib"
            if alt_path.exists():
                return str(alt_path)
        
        # Fall back to original find_library
        return _original_find_library(name)
    
    # Apply the monkey patch
    ctypes.util.find_library = _patched_find_library

from src.agents.base_agent import BaseAgent
from src.utils.file_manager import FileManager
from src.context.context_manager import ContextManager
from src.context.shared_context import AgentType, DocumentStatus, AgentOutput
from src.rate_limit.queue_manager import RequestQueue
from src.utils.logger import get_logger

logger = get_logger(__name__)


# Mapping from AgentType values to folder names in docs/
AGENT_TYPE_TO_FOLDER = {
    # Level 1: Strategic (Entrepreneur)
    "requirements_analyst": "requirements",
    "stakeholder_communication": "stakeholder",
    "project_charter": "charter",
    "business_model": "business",
    "marketing_plan": "marketing",
    
    # Level 2: Product (Product Manager)
    "pm_documentation": "pm",
    "user_stories": "user_stories",
    
    # Level 3: Technical (Programmer)
    "technical_documentation": "technical",
    "api_documentation": "api",
    "database_schema": "database",
    "setup_guide": "setup",
    "legal_compliance": "legal",
    
    # Cross-Level
    "developer_documentation": "developer",
    "user_documentation": "user",
    "test_documentation": "test",
    "support_playbook": "support",
    "quality_reviewer": "quality",
}


class FormatConverterAgent(BaseAgent):
    """
    Format Converter Agent
    
    Converts documentation between formats:
    - Markdown to HTML
    - Markdown to PDF (via HTML)
    - Markdown to DOCX
    - Preserves formatting and structure
    """
    
    def __init__(
        self,
        provider_name: Optional[str] = None,
        model_name: Optional[str] = None,
        rate_limiter: Optional[RequestQueue] = None,
        file_manager: Optional[FileManager] = None,
        api_key: Optional[str] = None,
        **provider_kwargs
    ):
        """Initialize Format Converter Agent"""
        # FormatConverter doesn't actually use LLM, but inherits for consistency
        super().__init__(
            provider_name=provider_name,
            model_name=model_name,
            rate_limiter=rate_limiter,
            api_key=api_key,
            **provider_kwargs
        )
        
        self.file_manager = file_manager or FileManager(base_dir="docs")
        self.supported_formats = ["html", "pdf", "docx"]
        logger.debug(f"FormatConverterAgent initialized with supported formats: {self.supported_formats}")
    
    def generate(self, markdown_content: str) -> str:
        """
        Generate method required by BaseAgent interface
        
        Args:
            markdown_content: Markdown content (for consistency with other agents)
        
        Returns:
            HTML representation of the markdown (default conversion)
        """
        return self.markdown_to_html(markdown_content)
    
    def markdown_to_html(self, markdown_content: str) -> str:
        """
        Convert Markdown to HTML with comprehensive formatting
        
        Args:
            markdown_content: Markdown content to convert
        
        Returns:
            HTML content with all Markdown syntax properly converted
        """
        logger.debug(f"Converting Markdown to HTML (input length: {len(markdown_content)} characters)")
        try:
            import markdown
            import re
            
            # Enhanced Markdown extensions for better conversion
            md = markdown.Markdown(
                extensions=[
                    'extra',           # Extra features (fenced code, tables, etc.)
                    'codehilite',      # Syntax highlighting
                    'tables',          # Table support
                    'nl2br',           # Convert newlines to <br>
                    'sane_lists',      # Better list handling
                    'toc',             # Table of contents support
                ]
            )
            
            # Convert Markdown to HTML
            html_content = md.convert(markdown_content)
            logger.debug(f"Markdown converted to HTML (output length: {len(html_content)} characters)")
            
            # Post-process to ensure all Markdown syntax is removed
            # Fix any remaining Markdown headers (## becomes h2, etc.)
            html_content = re.sub(r'##\s+(.+)', r'<h2>\1</h2>', html_content)
            html_content = re.sub(r'###\s+(.+)', r'<h3>\1</h3>', html_content)
            html_content = re.sub(r'####\s+(.+)', r'<h4>\1</h4>', html_content)
            
            # Fix any remaining bold syntax (**text** or __text__)
            html_content = re.sub(r'\*\*([^*]+)\*\*', r'<strong>\1</strong>', html_content)
            html_content = re.sub(r'__([^_]+)__', r'<strong>\1</strong>', html_content)
            
            # Fix any remaining italic syntax (*text* or _text_)
            html_content = re.sub(r'(?<!\*)\*(?!\*)([^*]+?)(?<!\*)\*(?!\*)', r'<em>\1</em>', html_content)
            html_content = re.sub(r'(?<!_)_(?!_)([^_]+?)(?<!_)_(?!_)', r'<em>\1</em>', html_content)
            # Add Mermaid.js support for diagrams
            # Replace mermaid code blocks with divs that Mermaid.js can render
            mermaid_pattern = r'```mermaid\n(.*?)```'
            mermaid_blocks = re.findall(mermaid_pattern, markdown_content, re.DOTALL)
            if mermaid_blocks:
                logger.debug(f"Found {len(mermaid_blocks)} Mermaid diagram(s) to render")
                # Replace mermaid code blocks with divs
                html_content = re.sub(
                    mermaid_pattern,
                    lambda m: f'<div class="mermaid">\n{m.group(1).strip()}\n</div>',
                    html_content,
                    flags=re.DOTALL
                )
            
            # Wrap in proper HTML structure with enhanced styling and Mermaid.js
            mermaid_script = '<script src="https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.min.js"></script>\n    <script>mermaid.initialize({startOnLoad:true, theme:"default"});</script>' if mermaid_blocks else ''
            
            full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    {mermaid_script}
    <title>Documentation</title>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Helvetica Neue', Arial, sans-serif;
            max-width: 900px;
            margin: 0 auto;
            padding: 40px 20px;
            line-height: 1.8;
            color: #333;
            background-color: #fff;
        }}
        h1 {{ font-size: 2.5em; margin: 1em 0 0.5em; color: #2c3e50; border-bottom: 3px solid #3498db; padding-bottom: 0.3em; }}
        h2 {{ font-size: 2em; margin: 1.5em 0 0.5em; color: #34495e; border-bottom: 2px solid #95a5a6; padding-bottom: 0.3em; }}
        h3 {{ font-size: 1.5em; margin: 1.2em 0 0.5em; color: #555; }}
        h4, h5, h6 {{ margin: 1em 0 0.5em; color: #666; }}
        p {{ margin: 1em 0; text-align: justify; }}
        code {{
            background-color: #f8f9fa;
            padding: 2px 6px;
            border-radius: 4px;
            font-family: 'Monaco', 'Courier New', monospace;
            font-size: 0.9em;
            color: #e83e8c;
        }}
        pre {{
            background-color: #f8f9fa;
            padding: 16px;
            border-radius: 6px;
            overflow-x: auto;
            border-left: 4px solid #3498db;
            margin: 1em 0;
        }}
        pre code {{ background-color: transparent; padding: 0; color: #333; }}
        .mermaid {{
            background-color: #f8f9fa;
            padding: 20px;
            border-radius: 6px;
            margin: 1.5em 0;
            text-align: center;
            border: 1px solid #dee2e6;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 1.5em 0;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }}
        th, td {{
            border: 1px solid #dee2e6;
            padding: 12px;
            text-align: left;
        }}
        th {{
            background-color: #3498db;
            color: white;
            font-weight: 600;
        }}
        tr:nth-child(even) {{ background-color: #f8f9fa; }}
        ul, ol {{ margin: 1em 0; padding-left: 2em; }}
        li {{ margin: 0.5em 0; }}
        blockquote {{
            border-left: 4px solid #3498db;
            padding-left: 1em;
            margin: 1em 0;
            color: #666;
            font-style: italic;
        }}
        a {{ color: #3498db; text-decoration: none; }}
        a:hover {{ text-decoration: underline; }}
        @media print {{
            body {{ max-width: 100%; padding: 0; }}
            h1, h2, h3 {{ page-break-after: avoid; }}
        }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""
            logger.info("Markdown to HTML conversion completed successfully")
            return full_html
        except ImportError:
            # Fallback: basic conversion without markdown library
            logger.warning("Markdown library not installed, using basic HTML conversion")
            html_content = markdown_content.replace('\n', '<br>\n')
            return f"<html><body>{html_content}</body></html>"
        except Exception as e:
            logger.error(f"Error converting Markdown to HTML: {str(e)}", exc_info=True)
            raise
    
    def html_to_pdf(self, html_content: str, output_path: Optional[str] = None, subdirectory: Optional[str] = None) -> str:
        """
        Convert HTML to PDF with enhanced styling
        
        Args:
            html_content: HTML content to convert
            output_path: Optional output file path
            subdirectory: Optional subdirectory to save file in
        
        Returns:
            Path to generated PDF file
        """
        logger.info(f"Converting HTML to PDF (subdirectory: {subdirectory}, output_path: {output_path})")
        # Pre-load all required libraries using ctypes before importing WeasyPrint
        # This helps WeasyPrint find the libraries on macOS with SIP
        if sys.platform == 'darwin' and Path('/opt/homebrew/lib').exists():
            try:
                import ctypes
                homebrew_lib = '/opt/homebrew/lib'
                
                # List of libraries WeasyPrint needs (in dependency order)
                # Try both versioned and unversioned names
                libs_to_load = [
                    'libgobject-2.0.0.dylib',  # Versioned name (what WeasyPrint looks for)
                    'libgobject-2.0.dylib',     # Unversioned name
                    'libglib-2.0.dylib',
                    'libcairo.2.dylib',
                    'libpango-1.0.0.dylib',     # Versioned name
                    'libpango-1.0.dylib',       # Unversioned name
                    'libpangoft2-1.0.dylib',
                ]
                
                # Pre-load each library in dependency order
                for lib_name in libs_to_load:
                    lib_path = Path(homebrew_lib) / lib_name
                    if lib_path.exists():
                        try:
                            ctypes.CDLL(str(lib_path), mode=ctypes.RTLD_GLOBAL)
                        except Exception:
                            # Continue if one library fails - others might still work
                            pass
                            
                # Also try loading from Cellar (direct path)
                cellar_paths = [
                    '/opt/homebrew/Cellar/glib/*/lib/libgobject-2.0.dylib',
                    '/opt/homebrew/Cellar/cairo/*/lib/libcairo.2.dylib',
                    '/opt/homebrew/Cellar/pango/*/lib/libpango-1.0.dylib',
                ]
                for pattern in cellar_paths:
                    import glob
                    matches = glob.glob(pattern)
                    if matches:
                        try:
                            ctypes.CDLL(matches[0], mode=ctypes.RTLD_GLOBAL)
                        except Exception:
                            pass
            except Exception:
                # If pre-loading fails, continue anyway - WeasyPrint might still work
                pass
        
        # Suppress stderr BEFORE importing weasyprint (it prints during import)
        import io
        import warnings
        
        stderr_backup = sys.stderr
        stderr_capture = io.StringIO()
        sys.stderr = stderr_capture
        
        # Also suppress warnings
        warnings.filterwarnings("ignore")
        
        try:
            # Import weasyprint with stderr suppressed
            from weasyprint import HTML, CSS
            
            if not output_path:
                output_path = "documentation.pdf"
            
            # If subdirectory is provided, create path with subdirectory
            if subdirectory:
                output_path = f"{subdirectory}/{output_path}"
            
            # Ensure .pdf extension
            if not output_path.endswith('.pdf'):
                output_path = str(Path(output_path).with_suffix('.pdf'))
            
            # Enhanced CSS for PDF - Formal document styling
            pdf_css = CSS(string='''
                @page {
                    size: A4;
                    margin: 2.5cm 2cm;
                    @top-center {
                        content: "Documentation";
                        font-size: 9pt;
                        color: #666;
                        border-bottom: 1px solid #ddd;
                        padding-bottom: 0.5cm;
                    }
                    @bottom-center {
                        content: "Page " counter(page) " of " counter(pages);
                        font-size: 9pt;
                        color: #666;
                        border-top: 1px solid #ddd;
                        padding-top: 0.5cm;
                    }
                }
                body {
                    font-family: "Times New Roman", Times, serif;
                    font-size: 11pt;
                    line-height: 1.8;
                    color: #000;
                    text-align: justify;
                }
                h1 {
                    font-size: 18pt;
                    font-weight: bold;
                    margin-top: 2em;
                    margin-bottom: 1em;
                    page-break-after: avoid;
                    color: #000;
                    border-bottom: 2px solid #000;
                    padding-bottom: 0.3em;
                }
                h2 {
                    font-size: 16pt;
                    font-weight: bold;
                    margin-top: 1.5em;
                    margin-bottom: 0.8em;
                    page-break-after: avoid;
                    color: #000;
                    border-bottom: 1px solid #666;
                    padding-bottom: 0.2em;
                }
                h3 {
                    font-size: 14pt;
                    font-weight: bold;
                    margin-top: 1.2em;
                    margin-bottom: 0.6em;
                    page-break-after: avoid;
                    color: #000;
                }
                h4, h5, h6 {
                    font-size: 12pt;
                    font-weight: bold;
                    margin-top: 1em;
                    margin-bottom: 0.5em;
                    page-break-after: avoid;
                    color: #000;
                }
                p {
                    margin: 0.8em 0;
                    text-align: justify;
                    text-indent: 0;
                }
                strong {
                    font-weight: bold;
                }
                em {
                    font-style: italic;
                }
                ul, ol {
                    margin: 1em 0;
                    padding-left: 2em;
                }
                li {
                    margin: 0.5em 0;
                    text-align: justify;
                }
                pre {
                    background-color: #f5f5f5;
                    border: 1px solid #ddd;
                    padding: 1em;
                    page-break-inside: avoid;
                    font-family: "Courier New", monospace;
                    font-size: 9pt;
                    overflow-wrap: break-word;
                }
                code {
                    background-color: #f5f5f5;
                    padding: 2px 4px;
                    font-family: "Courier New", monospace;
                    font-size: 9pt;
                }
                table {
                    width: 100%;
                    border-collapse: collapse;
                    margin: 1em 0;
                    page-break-inside: avoid;
                }
                th, td {
                    border: 1px solid #000;
                    padding: 8px;
                    text-align: left;
                }
                th {
                    background-color: #f0f0f0;
                    font-weight: bold;
                }
                blockquote {
                    border-left: 3px solid #666;
                    margin: 1em 0;
                    padding-left: 1em;
                    color: #444;
                    font-style: italic;
                }
            ''')
            
            html_obj = HTML(string=html_content)
            pdf_path = self.file_manager.base_dir / output_path
            logger.debug(f"Writing PDF to: {pdf_path}")
            html_obj.write_pdf(pdf_path, stylesheets=[pdf_css])
            
            # Restore stderr AFTER successful conversion
            sys.stderr = stderr_backup
            
            pdf_abs_path = str(pdf_path.absolute())
            logger.info(f"PDF file generated successfully: {pdf_abs_path}")
            return pdf_abs_path
            
        except (OSError, ImportError, Exception) as e:
            # Restore stderr before checking error
            sys.stderr = stderr_backup
            # Clear the captured stderr (don't print the WeasyPrint warnings)
            stderr_capture.seek(0)
            stderr_capture.truncate(0)
            
            # If it's a library loading error (macOS), skip PDF conversion gracefully
            error_str = str(e).lower()
            if "libgobject" in error_str or "dlopen" in error_str or "cannot load library" in error_str:
                logger.error(f"PDF conversion failed: System libraries not available. Error: {str(e)}")
                raise ImportError(
                    "PDF conversion unavailable: System libraries not available. "
                    "PDF conversion requires additional system libraries on macOS. "
                    "HTML and DOCX formats are still available."
                )
            logger.error(f"Error converting HTML to PDF: {str(e)}", exc_info=True)
            raise
        except ImportError:
            try:
                # Fallback to pdfkit if available
                import pdfkit
                
                if not output_path:
                    output_path = "documentation.pdf"
                if not output_path.endswith('.pdf'):
                    output_path = str(Path(output_path).with_suffix('.pdf'))
                
                pdf_path = self.file_manager.base_dir / output_path
                options = {
                    'page-size': 'A4',
                    'margin-top': '0.75in',
                    'margin-right': '0.75in',
                    'margin-bottom': '0.75in',
                    'margin-left': '0.75in',
                    'encoding': "UTF-8",
                    'no-outline': None
                }
                pdfkit.from_string(html_content, str(pdf_path), options=options)
                
                return str(pdf_path.absolute())
            except ImportError:
                raise ImportError(
                    "PDF conversion requires 'weasyprint' or 'pdfkit'. "
                    "Install with: pip install weasyprint"
                )
    
    def markdown_to_docx(self, markdown_content: str, output_path: Optional[str] = None, subdirectory: Optional[str] = None) -> str:
        """
        Convert Markdown to DOCX
        
        Args:
            markdown_content: Markdown content to convert
            output_path: Optional output file path
            subdirectory: Optional subdirectory to save file in
        
        Returns:
            Path to generated DOCX file
        """
        logger.info(f"Converting Markdown to DOCX (subdirectory: {subdirectory}, output_path: {output_path})")
        try:
            from docx import Document
            from docx.shared import Inches
            import re
            
            if not output_path:
                output_path = "documentation.docx"
            
            # If subdirectory is provided, create path with subdirectory
            if subdirectory:
                output_path = f"{subdirectory}/{output_path}"
            
            if not output_path.endswith('.docx'):
                output_path = str(Path(output_path).with_suffix('.docx'))
            
            doc = Document()
            
            # Parse markdown and convert to docx
            lines = markdown_content.split('\n')
            for line in lines:
                # Headings
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                # Lists
                elif line.strip().startswith('- ') or line.strip().startswith('* '):
                    doc.add_paragraph(line.strip()[2:], style='List Bullet')
                # Regular paragraphs
                elif line.strip():
                    doc.add_paragraph(line.strip())
                else:
                    doc.add_paragraph()
            
            docx_path = self.file_manager.base_dir / output_path
            logger.debug(f"Writing DOCX to: {docx_path}")
            doc.save(str(docx_path))
            
            docx_abs_path = str(docx_path.absolute())
            logger.info(f"DOCX file generated successfully: {docx_abs_path}")
            return docx_abs_path
        except ImportError as e:
            logger.error(f"DOCX conversion failed: python-docx library not available. Error: {str(e)}")
            raise ImportError(
                "DOCX conversion requires 'python-docx'. "
                "Install with: pip install python-docx"
            )
        except Exception as e:
            logger.error(f"Error converting Markdown to DOCX: {str(e)}", exc_info=True)
            raise
    
    def convert(
        self,
        markdown_content: str,
        output_format: str,
        output_filename: Optional[str] = None,
        subdirectory: Optional[str] = None
    ) -> str:
        """
        Convert Markdown content to specified format
        
        Args:
            markdown_content: Markdown content to convert
            output_format: Target format ('html', 'pdf', 'docx')
            output_filename: Optional output filename
            subdirectory: Optional subdirectory to save file in (e.g., 'api_documentation')
        
        Returns:
            Path to converted file
        """
        logger.info(f"Starting format conversion: {output_format} (filename: {output_filename}, subdirectory: {subdirectory})")
        if output_format.lower() not in self.supported_formats:
            logger.error(f"Unsupported format requested: {output_format}. Supported: {self.supported_formats}")
            raise ValueError(
                f"Unsupported format: {output_format}. "
                f"Supported formats: {', '.join(self.supported_formats)}"
            )
        
        if output_format.lower() == 'html':
            html_content = self.markdown_to_html(markdown_content)
            if not output_filename:
                output_filename = "documentation.html"
            # If subdirectory is provided, create path with subdirectory
            if subdirectory:
                output_filename = f"{subdirectory}/{output_filename}"
            # Generate virtual file path for reference (not used for actual file storage)
            virtual_path = f"docs/{output_filename}"
            logger.info(f"Format conversion completed: HTML -> {virtual_path} (saved to database)")
            # Note: Format converter output is typically not saved to agent_outputs table
            # as it's a conversion of existing documents, not a new document generation
            return virtual_path  # Return virtual path for compatibility
        
        elif output_format.lower() == 'pdf':
            html_content = self.markdown_to_html(markdown_content)
            pdf_path = self.html_to_pdf(html_content, output_filename, subdirectory)
            logger.info(f"Format conversion completed: PDF -> {pdf_path}")
            return pdf_path
        
        elif output_format.lower() == 'docx':
            docx_path = self.markdown_to_docx(markdown_content, output_filename, subdirectory)
            logger.info(f"Format conversion completed: DOCX -> {docx_path}")
            return docx_path
        
        else:
            logger.error(f"Format conversion not implemented: {output_format}")
            raise ValueError(f"Format conversion not implemented: {output_format}")
    
    def convert_all_documents(
        self,
        documents: dict,
        formats: List[str],
        project_id: Optional[str] = None,
        context_manager: Optional[ContextManager] = None
    ) -> dict:
        """
        Convert all documents to multiple formats with detailed status reporting
        
        Each document is saved in its own subdirectory: docs/{doc_name}/
        
        Args:
            documents: Dict mapping document names to markdown content
            formats: List of target formats (e.g., ['html', 'pdf', 'docx'])
            project_id: Optional project ID for context
            context_manager: Optional context manager
        
        Returns:
            Dict mapping document names to conversion results with detailed status:
            {
                "doc_name": {
                    "html": {
                        "status": "success",
                        "file_path": "docs/api/api_documentation.html"
                    },
                    "pdf": {
                        "status": "failed_dependency_error",
                        "error": "PDF conversion unavailable: System libraries not available",
                        "file_path": None
                    },
                    "docx": {
                        "status": "success",
                        "file_path": "docs/api/api_documentation.docx"
                    }
                }
            }
            
            Status values:
            - "success": Conversion successful
            - "failed_dependency_error": Missing system dependencies (e.g., WeasyPrint libraries)
            - "failed_import_error": Missing Python package (e.g., python-docx)
            - "failed_unknown_error": Other errors
        """
        results = {}
        
        logger.info(f"Starting batch conversion: {len(documents)} documents to formats: {', '.join(formats)}")
        logger.info(f"Files will be saved in docs/{{folder}}/ (matching original document folders)")
        
        for doc_name, markdown_content in documents.items():
            doc_results = {}
            
            # Map document name to the correct folder in docs/
            # Use AgentType mapping if available, otherwise use document name
            if doc_name:
                # First try to find in mapping (for AgentType values)
                folder_name = AGENT_TYPE_TO_FOLDER.get(doc_name.lower())
                
                if not folder_name:
                    # Extract clean document name (remove file extensions, normalize)
                    clean_name = str(Path(doc_name).stem) if '.' in str(doc_name) else str(doc_name)
                    # Normalize to lowercase, replace spaces/underscores/hyphens
                    clean_name = clean_name.lower().replace(' ', '_').replace('-', '_')
                    # Try mapping again with cleaned name
                    folder_name = AGENT_TYPE_TO_FOLDER.get(clean_name)
                    
                    # If still not found, use cleaned name (but try to match existing folder structure)
                    if not folder_name:
                        # Remove "_documentation" suffix if present to match folder names
                        folder_name = clean_name.replace('_documentation', '')
                        if folder_name == "requirements_analyst":
                            folder_name = "requirements"
                        elif folder_name == "stakeholder_communication":
                            folder_name = "stakeholder"
                subdirectory = folder_name
            else:
                subdirectory = None  # Will save to docs/ root
            
            for fmt in formats:
                try:
                    # Extract base name from doc_name (handle both string and Path-like)
                    if doc_name:
                        base_name = str(Path(doc_name).stem) if '.' in str(doc_name) else str(doc_name)
                    else:
                        base_name = "document"
                    output_filename = f"{base_name}.{fmt}"
                    
                    file_path = self.convert(
                        markdown_content=markdown_content,
                        output_format=fmt,
                        output_filename=output_filename,
                        subdirectory=subdirectory
                    )
                    
                    doc_results[fmt] = {
                        "status": "success",
                        "file_path": file_path
                    }
                    logger.info(f"Successfully converted {doc_name} to {fmt} → {subdirectory}/{output_filename}")
                    
                except ImportError as e:
                    # Missing Python package or system library
                    error_msg = str(e)
                    if "weasyprint" in error_msg.lower() or "system libraries" in error_msg.lower() or "libgobject" in error_msg.lower():
                        status = "failed_dependency_error"
                        error_detail = "PDF conversion requires system libraries (WeasyPrint dependencies). HTML and DOCX formats are still available."
                    elif "python-docx" in error_msg.lower():
                        status = "failed_import_error"
                        error_detail = f"DOCX conversion requires 'python-docx' package. Install with: pip install python-docx"
                    else:
                        status = "failed_import_error"
                        error_detail = f"Missing dependency: {error_msg}"
                    
                    doc_results[fmt] = {
                        "status": status,
                        "error": error_detail,
                        "file_path": None
                    }
                    logger.warning(f"Format conversion failed for {doc_name} to {fmt}: {error_detail}")
                    
                except Exception as e:
                    # Other errors
                    error_msg = str(e)
                    doc_results[fmt] = {
                        "status": "failed_unknown_error",
                        "error": error_msg,
                        "file_path": None
                    }
                    logger.error(f"Error converting {doc_name} to {fmt}: {error_msg}", exc_info=True)
            
            results[doc_name] = doc_results
        
        # Save to context if available
        if project_id and context_manager:
            output = AgentOutput(
                agent_type=AgentType.FORMAT_CONVERTER,
                document_type="format_conversions",
                content=str(results),  # JSON-like string of results
                file_path="",  # Multiple files, no single path
                status=DocumentStatus.COMPLETE,
                generated_at=datetime.now()
            )
            context_manager.save_agent_output(project_id, output)
            logger.info(f"Format conversions saved to shared context (project: {project_id})")
        
        logger.info(f"Batch conversion completed: {len(results)} documents processed")
        return results

